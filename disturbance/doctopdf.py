import os
from django.conf import settings
# from reportlab.lib import enums
# from reportlab.lib.pagesizes import A4
# from reportlab.platypus import BaseDocTemplate, PageTemplate, Frame, Paragraph, Table, TableStyle
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.lib.utils import ImageReader
# from mooring.models import Booking
import re
from docx import Document
from docxtpl import DocxTemplate
from datetime import datetime, timedelta, date


def docx_replace_regex(doc_obj, regex, replace, key, bold_font=False, italic_font=False, underline_font=False):
    for p in doc_obj.paragraphs:
        # p.text = regex.sub(replace, p.text)
        style = p.style
        if key in p.text:
            #    print (p.text)
            p.bold = True
            p.text = regex.sub(replace, '')
            p.add_run(replace).bold = bold_font
            # p.style.font.bold = True
        if regex.search(p.text):
            inline = p.runs
            # print (inline)
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            # print (row)
            for cell in row.cells:
                docx_replace_regex(cell, regex, replace, key, bold_font, italic_font, underline_font)


def create_annual_admission_letter(booking):
    print ("Letter File")
    confirmation_doc = None
    if booking.annual_booking_period_group.letter:
        print (booking.annual_booking_period_group.letter.path)
        confirmation_doc = booking.annual_booking_period_group.letter.path
    # confirmation_doc = settings.BASE_DIR+"/mooring/templates/doc/AnnualAdmissionStickerLetter.docx"
    temp_directory = settings.BASE_DIR + "/tmp/"

    doc = DocxTemplate(confirmation_doc)
    address = ''
    if len(booking.details.get('postal_address_line_2', '')) > 0:
        address = '{}, {}'.format(booking.details.get('postal_address_line_1', ''),
                                  booking.details.get('postal_address_line_2', ''))
    else:
        address = '{}'.format(booking.details.get('postal_address_line_1', ''))
    bookingdate = booking.created + timedelta(hours=8)
    todaydate = datetime.utcnow() + timedelta(hours=8)
    stickercreated = ''
    if booking.sticker_created:
        sc = booking.sticker_created + timedelta(hours=8)
        stickercreated = sc.strftime('%d %B %Y')

    context = {
        'customername': '{} {}'.format(booking.details.get('first_name', ''), booking.details.get('last_name', '')),
        'customeraddress': address, "customersuburb": booking.details.get('suburb', ''),
        "customerstate": booking.details.get('state', ''), 'customerpostcode': booking.details.get('post_code', ''),
        'bookingyear': '{}/{}'.format(booking.annual_booking_period_group.start_time.strftime('%Y'),
                                      booking.annual_booking_period_group.finish_time.strftime('%y')),
        'admissionsexpiry': booking.annual_booking_period_group.finish_time.strftime('%d %B %Y'),
        'vessel': booking.details.get('vessel_rego', ''), 'customerfirstname': booking.details.get('first_name', ''),
        'bookingdate': booking.created.strftime('%d %B %Y'), 'todaydate': todaydate.strftime('%d %B %Y'),
        'stickercreated': stickercreated}
    doc.render(context)

    try:
        os.stat(temp_directory)
    except:
        os.mkdir(temp_directory)
    new_doc_file = temp_directory + 'booking_confirmation_' + str(booking.id) + '.docx'
    new_pdf_file = temp_directory + 'booking_confirmation_' + str(booking.id) + '.pdf'
    doc.save(new_doc_file)
    os.system("libreoffice --headless --convert-to pdf " + new_doc_file + " --outdir " + temp_directory)

    confirmation_buffer = None
    with open(new_pdf_file, 'rb') as f:
        confirmation_buffer = f.read()
    os.remove(new_doc_file)
    os.remove(new_pdf_file)
    return confirmation_buffer


def create_annual_admission_letter_old(booking):
    print
    confirmation_doc = settings.BASE_DIR + "/mooring/templates/doc/AnnualAdmissionStickerLetter.docx"
    temp_directory = settings.BASE_DIR + "/tmp/"

    doc = Document(confirmation_doc)

    # if booking.first_campsite_list:
    #    campsites = []
    #    if booking.campground.site_type == 0:
    #        for item in booking.first_campsite_list:
    #            campsites.append(item.name if item else "")
    #    elif booking.campground.site_type == 1 or 2:
    #        for item in booking.first_campsite_list:
    #            campsites.append(item.type.split(':', 1)[0] if item else "")
    #    campsite = ', '.join(campsites)
    #    result = {x: campsites.count(x) for x in campsites}
    #    for key, value in result.items():
    #        campsite = ', '.join(['%sx %s' % (value, key) for (key, value) in result.items()])

    # vehicle_data = ""
    # if booking.vehicle_payment_status:
    #    for r in booking.vehicle_payment_status:
    #        data = r['Type']+"\t"+r['Rego']
    #        if r.get('Paid') is not None:
    #            if r['Paid'] == 'Yes':
    #               data = data+'\tEntry fee paid'
    #            elif r['Paid'] == 'No':
    #               data = data+'\tUnpaid'
    #            elif r['Paid'] == 'pass_required':
    #               data = data+'\tPark Pass Required'
    #        vehicle_data = vehicle_data + data+"\n"

    # key = "{{campname}}"
    # docx_replace_regex(doc, re.compile(r''+key) , '{}, {}'.format('', ''), key, True)
    key = "{{customername}}"
    docx_replace_regex(doc, re.compile(r'' + key),
                       '{} {}'.format(booking.details.get('first_name', ''), booking.details.get('last_name', '')), key,
                       False)
    key = "{{customeraddress}}"
    if len(booking.details.get('postal_address_line_2', '')) > 0:
        docx_replace_regex(doc, re.compile(r'' + key), '{}\n{}'.format(booking.details.get('postal_address_line_1', ''),
                                                                       booking.details.get('postal_address_line_2',
                                                                                           '')), key, False)
    else:
        docx_replace_regex(doc, re.compile(r'' + key), '{}'.format(booking.details.get('postal_address_line_1', '')),
                           key, False)
    # key = "{{customersuburb}"
    # docx_replace_regex(doc, re.compile(r''+key) , '{}'.format(booking.details.get('suburb', '')), key,False)
    key = "{{customerstate}"
    docx_replace_regex(doc, re.compile(r'' + key), '{}'.format(booking.details.get('state', '')), key, False)
    key = "{{customerpostcode}"
    docx_replace_regex(doc, re.compile(r'' + key), '{}'.format(booking.details.get('post_code', '')), key, False)
    # key = "{{admissionsexpiry}"
    # docx_replace_regex(doc, re.compile(r''+key) , '{}'.format(booking.details.get('post_code', '')), key,False)

    try:
        os.stat(temp_directory)
    except:
        os.mkdir(temp_directory)
    new_doc_file = temp_directory + 'booking_confirmation_' + str(booking.id) + '.docx'
    new_pdf_file = temp_directory + 'booking_confirmation_' + str(booking.id) + '.pdf'
    doc.save(new_doc_file)
    os.system("libreoffice --headless --convert-to pdf " + new_doc_file + " --outdir " + temp_directory)

    confirmation_buffer = None
    with open(new_pdf_file, 'rb') as f:
        confirmation_buffer = f.read()
    os.remove(new_doc_file)
    os.remove(new_pdf_file)
    return confirmation_buffer